#!/usr/bin/env python
"""Build a Word .docx from the COMAVI manuscript markdown.

Handles: ATX headings, pipe tables, ordered/unordered lists (including
wrapped continuation lines), block quotes, fenced code, inline emphasis
(**bold**, *italic*, `code`), and figure embeds written as
`![caption]({{artifact:art_<uuid>}})`.

Figure artifacts are resolved through the host artifact store when
available, otherwise from a local figures/ directory.

Usage:
    python scripts/build_manuscript_docx.py docs/COMAVI_manuscript_v22.md \
           docs/COMAVI_manuscript_v22.docx
"""
import json
import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

TEXT_WIDTH_IN = 6.5

IMG = re.compile(r"^!\[([^\]]*)\]\(\{\{artifact:art_([0-9a-f\-]+)\}\}\)\s*$")
INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def resolve_artifact(aid, fallback_dir=None):
    """Return a filesystem path for an artifact id.

    Resolution order is repo-first: figures/manifest.json maps each artifact
    id embedded in the manuscript to its committed PNG, so a checkout builds
    the same document without needing the artifact store. The store is
    consulted only for ids the manifest does not cover.

    Raises rather than returning None. An earlier version fell back to "the
    first PNG in the directory", which embedded whichever file happened to
    sort first under all eight figure ids; when the store was also
    unreachable it returned None and the build emitted a figure-less
    manuscript with exit status 0. A missing figure must stop the build.
    """
    fig_dir = Path(fallback_dir) if fallback_dir else Path(__file__).resolve().parent.parent / "figures"
    manifest_path = fig_dir / "manifest.json"
    if manifest_path.exists():
        manifest = json.loads(manifest_path.read_text())
        if aid in manifest:
            p = fig_dir / manifest[aid]
            if not p.exists():
                raise FileNotFoundError(
                    f"manifest maps {aid} to {manifest[aid]}, which is absent from {fig_dir}")
            return str(p)
    try:
        host  # noqa: F821  (injected in kernel contexts)
    except NameError:
        pass
    else:
        try:
            return host.artifact_path(aid)  # noqa: F821
        except Exception:
            pass
    raise KeyError(
        f"figure artifact {aid} is in neither {manifest_path} nor the artifact store")


def add_md_runs(par, text):
    """Add text to a paragraph, honouring inline markdown emphasis."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            par.add_run(tok)


def coalesce_lists(lines):
    """Join wrapped continuation lines into their parent list item.

    A hard-wrapped markdown list item spans several physical lines; emitting
    each as its own paragraph splits items mid-sentence.
    """
    out, buf = [], None
    item = re.compile(r"^\s*(?:[-*+]|\d+\.)\s+")
    for ln in lines:
        if item.match(ln):
            if buf is not None:
                out.append(buf)
            buf = ln.rstrip()
        elif buf is not None and ln.strip() and not ln.startswith(("|", "#", ">", "!")):
            buf += " " + ln.strip()
        else:
            if buf is not None:
                out.append(buf)
                buf = None
            out.append(ln)
    if buf is not None:
        out.append(buf)
    return out


def build_docx(md, outpath, fig_dir=None):
    doc = docx.Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)

    lines = coalesce_lists(md.split("\n"))
    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]

        if not ln.strip() or ln.strip() == "---":
            i += 1
            continue

        mo = IMG.match(ln.strip())
        if mo:
            cap, aid = mo.group(1), mo.group(2)
            path = resolve_artifact(aid, fig_dir)
            if path:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(TEXT_WIDTH_IN))
            else:
                doc.add_paragraph(f"[figure unavailable: {cap}]")
            i += 1
            continue

        if ln.startswith("#"):
            lvl = len(ln) - len(ln.lstrip("#"))
            doc.add_heading(ln.lstrip("# ").strip(), level=min(lvl, 4))
            i += 1
            continue

        if ln.startswith("|"):
            block = []
            while i < n and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in block
                if not re.match(r"^\|[\s:\-|]+\|?$", r.strip())
            ]
            if rows:
                t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        c = t.cell(ri, ci)
                        c.text = ""
                        add_md_runs(c.paragraphs[0], cell)
                        if ri == 0:
                            for r_ in c.paragraphs[0].runs:
                                r_.bold = True
                        for r_ in c.paragraphs[0].runs:
                            r_.font.size = Pt(9)
            continue

        if ln.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_md_runs(p, ln.lstrip("> ").strip())
            for r_ in p.runs:
                r_.italic = True
                r_.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        if ln.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            r_ = p.add_run("\n".join(code))
            r_.font.name = "Consolas"
            r_.font.size = Pt(8.5)
            continue

        mo = re.match(r"^\s*(?:[-*+]|(\d+)\.)\s+(.*)$", ln)
        if mo:
            style = "List Number" if mo.group(1) else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_md_runs(p, mo.group(2))
            i += 1
            continue

        # plain paragraph: gather wrapped lines
        para = []
        while i < n and lines[i].strip() and not lines[i].startswith(
            ("#", "|", ">", "!", "```", "---")
        ) and not re.match(r"^\s*(?:[-*+]|\d+\.)\s+", lines[i]):
            para.append(lines[i].strip())
            i += 1
        if para:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            add_md_runs(p, " ".join(para))

    doc.save(outpath)
    return outpath


if __name__ == "__main__":
    src, dst = sys.argv[1], sys.argv[2]
    fig = sys.argv[3] if len(sys.argv) > 3 else None
    build_docx(Path(src).read_text(), dst, fig)
    print(f"wrote {dst}")
